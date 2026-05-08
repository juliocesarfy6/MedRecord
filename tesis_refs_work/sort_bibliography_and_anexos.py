from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Pt


DOCX = Path(r"C:\Users\julio\Gestor_CHM\MedRecord\tesis_refs_work\Plantilla_de_tesis_ampliada.docx")


def remove_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def insert_paragraph_before(target, text="", style="normal"):
    new_p = OxmlElement("w:p")
    target._element.addprevious(new_p)
    p = Paragraph(new_p, target._parent)
    p.style = style
    if text:
        p.add_run(text)
    return p


def is_page_break_only(paragraph):
    text = paragraph.text.strip()
    if text:
        return False
    return bool(paragraph._element.xpath(".//w:br[@w:type='page']"))


doc = Document(DOCX)

bib = None
anexos = None
for p in doc.paragraphs:
    if p.text.strip().startswith("Bibliograf"):
        bib = p
    if p.text.strip() == "Anexos":
        anexos = p
        break

if bib is None or anexos is None:
    raise SystemExit("Bibliografía o Anexos no encontrados")

refs = []
collect = False
to_remove = []
for p in list(doc.paragraphs):
    if p._element is bib._element:
        collect = True
        continue
    if p._element is anexos._element:
        break
    if collect:
        if p.text.strip():
            refs.append(p.text.strip())
        to_remove.append(p)

for p in to_remove:
    remove_paragraph(p)

for ref in sorted(refs, key=lambda x: x.lower()):
    para = insert_paragraph_before(anexos, ref, "normal")
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.first_line_indent = Inches(-0.5)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.0
    for run in para.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10)

page_break = insert_paragraph_before(anexos)
page_break.add_run().add_break(WD_BREAK.PAGE)

doc.save(DOCX)
print(f"Sorted bibliography and separated Anexos in {DOCX}")
