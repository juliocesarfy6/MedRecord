from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.shared import Inches, Pt


DOCX = Path(r"C:\Users\julio\Gestor_CHM\MedRecord\tesis_refs_work\Plantilla_de_tesis_citas_actualizada.docx")


APPOINTMENT_REFS = [
    "Cayirli, T., & Veral, E. (2003). Outpatient scheduling in health care: A review of literature. Production and Operations Management, 12(4), 519-549. https://doi.org/10.1111/j.1937-5956.2003.tb00218.x",
    "Gupta, D., & Denton, B. (2008). Appointment scheduling in health care: Challenges and opportunities. IIE Transactions, 40(9), 800-819. https://doi.org/10.1080/07408170802165880",
    "Zhao, P., Yoo, I., Lavoie, J., Lavoie, B. J., & Simoes, E. (2017). Web-based medical appointment systems: A systematic review. Journal of Medical Internet Research, 19(4), e134. https://doi.org/10.2196/jmir.6747",
]


def iter_block_items(parent):
    body = parent.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def paragraph_after(paragraph, text="", style="normal"):
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    p.style = style
    if text:
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(12)
        if style.startswith("Heading"):
            run.bold = True
    p.paragraph_format.space_after = Pt(6)
    return p


def paragraph_before(paragraph, text="", style="normal"):
    new_p = OxmlElement("w:p")
    paragraph._element.addprevious(new_p)
    p = Paragraph(new_p, paragraph._parent)
    p.style = style
    if text:
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(12)
        if style.startswith("Heading"):
            run.bold = True
    p.paragraph_format.space_after = Pt(6)
    return p


def set_text(paragraph, text, style=None):
    if style:
        paragraph.style = style
    alignment = paragraph.alignment
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    paragraph.alignment = alignment


def find_para(doc, predicate):
    for paragraph in doc.paragraphs:
        if predicate(paragraph):
            return paragraph
    return None


def add_after(anchor, items):
    current = anchor
    for style, text in items:
        current = paragraph_after(current, text, style)
    return current


def set_cell_text(cell, text):
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)


def set_table_rows(table, rows):
    for row, (label, value) in zip(table.rows, rows):
        set_cell_text(row.cells[0], label)
        set_cell_text(row.cells[1], value)


def insert_table_after(element, template_table, rows):
    new_tbl = deepcopy(template_table._element)
    element.addnext(new_tbl)
    table = Table(new_tbl, template_table._parent)
    set_table_rows(table, rows)
    return table


def table_title(table):
    if len(table.rows) and len(table.rows[0].cells) > 1:
        return table.rows[0].cells[1].text.strip()
    return ""


def find_table(doc, title):
    for table in doc.tables:
        if table_title(table) == title:
            return table
    return None


def remove_bibliography_between(doc, bib, anexos):
    refs = []
    collecting = False
    to_remove = []
    for p in list(doc.paragraphs):
        if p._element is bib._element:
            collecting = True
            continue
        if p._element is anexos._element:
            break
        if collecting:
            if p.text.strip():
                refs.append(p.text.strip())
            to_remove.append(p)
    for p in to_remove:
        p._element.getparent().remove(p._element)
    return refs


doc = Document(DOCX)

# Main objective and selected introductory/problem statements now include appointment management.
replacements = {
    "Diseñar e implementar un sistema web centrado en el paciente para la gestión y control de acceso al historial médico electrónico, incorporando mecanismos de autenticación, autorización y trazabilidad alineados con principios de confidencialidad, integridad y disponibilidad de la información (ISO/IEC, 2022a; NIST, 2020).":
        "Diseñar e implementar un sistema web centrado en el paciente para la gestión de citas médicas y el control de acceso al historial médico electrónico, incorporando mecanismos de autenticación, autorización, trazabilidad y administración de disponibilidad médica alineados con principios de confidencialidad, integridad y disponibilidad de la información (ISO/IEC, 2022a; NIST, 2020; Zhao et al., 2017).",
    "Diseñar e implementar la arquitectura del sistema web, incluyendo el modelo de base de datos y el esquema de control de acceso centrado en el paciente mediante generación y validación de tokens temporales.":
        "Diseñar e implementar la arquitectura del sistema web, incluyendo el modelo de base de datos, el esquema de control de acceso centrado en el paciente mediante generación y validación de tokens temporales, y el módulo de gestión de citas médicas con disponibilidad, agendamiento, cancelación y seguimiento por estado.",
    "Validar el funcionamiento del sistema mediante pruebas técnicas y evaluación con usuarios potenciales, utilizando métricas de desempeño, disponibilidad de información y percepción de seguridad.":
        "Validar el funcionamiento del sistema mediante pruebas técnicas y evaluación con usuarios potenciales, utilizando métricas de desempeño, disponibilidad de información, percepción de seguridad y claridad en el flujo de agendamiento de citas.",
    "El problema radica en la inexistencia de un sistema web interoperable y seguro, centrado en el paciente, que permita centralizar el historial médico electrónico, gestionar autorizaciones explícitas de acceso y garantizar trazabilidad en la consulta de información clínica, afectando la continuidad, calidad y seguridad de la atención médica.":
        "El problema radica en la inexistencia de un sistema web interoperable y seguro, centrado en el paciente, que permita centralizar el historial médico electrónico, gestionar citas médicas, administrar autorizaciones explícitas de acceso y garantizar trazabilidad en la consulta de información clínica, afectando la continuidad, calidad y seguridad de la atención médica.",
}

for p in doc.paragraphs:
    if p.text.strip() in replacements:
        set_text(p, replacements[p.text.strip()])

# Add a theory subsection on appointment management before related works.
related = find_para(doc, lambda p: p.text.strip() == "2.2 Trabajos relacionados")
if related and not find_para(doc, lambda p: p.text.strip().startswith("2.1.7 Gestión de citas")):
    items = [
        ("Heading 3", "2.1.7 Gestión de citas médicas y disponibilidad"),
        ("normal", "La gestión de citas médicas constituye un componente operativo relevante dentro de los sistemas de información en salud, ya que funciona como punto de entrada para muchos servicios ambulatorios. Un sistema de citas no solo registra una fecha y una hora; también relaciona la demanda del paciente con la capacidad disponible del profesional de salud. La literatura sobre programación de citas señala que un buen esquema de agenda debe equilibrar tiempos de espera, utilización del personal médico, preferencias de los pacientes y restricciones del servicio (Cayirli & Veral, 2003; Gupta & Denton, 2008)."),
        ("normal", "En sistemas web, la gestión de citas adquiere un valor adicional porque permite que el paciente consulte médicos disponibles, seleccione horarios, registre el motivo de la atención y dé seguimiento al estado de sus solicitudes sin depender exclusivamente de procesos presenciales o telefónicos. Las revisiones sobre sistemas web de citas médicas reportan beneficios asociados con mayor participación del paciente, reducción de carga administrativa, disminución de tiempos de espera y mejora en la satisfacción, aunque también advierten retos relacionados con adopción, seguridad, flexibilidad y calidad de la información registrada (Zhao et al., 2017)."),
        ("normal", "Para esta tesis, la gestión de citas se integra como una extensión natural del historial médico electrónico. La cita representa el evento que antecede o acompaña la atención clínica, mientras que el historial médico conserva el resultado documental de dicha atención. Por ello, el sistema propuesto relaciona pacientes, médicos, disponibilidad, citas y registros clínicos dentro de una misma plataforma. Esta relación permite que el paciente no solo comparta su información mediante tokens, sino que también pueda solicitar atención con un médico validado y mantener un seguimiento organizado de sus citas pendientes, confirmadas, reprogramadas, canceladas o completadas."),
        ("normal", "El módulo de citas implementado considera reglas básicas de disponibilidad: el médico define días y horarios activos, el sistema genera espacios de atención de duración fija y se evita que dos citas activas ocupen el mismo intervalo. Además, cada cita conserva un estado operativo que permite conocer su avance dentro del flujo de atención. Esta lógica aporta orden al proceso de atención y complementa los mecanismos de seguridad, porque las operaciones de creación, cancelación, confirmación, reprogramación y cierre quedan vinculadas a usuarios autenticados y pueden registrarse en auditoría."),
    ]
    for style, text in items:
        paragraph_before(related, text, style)

# Add appointment-related requirements after functional requirements.
rf_anchor = find_para(doc, lambda p: p.text.startswith("Finalmente, el sistema debe registrar en un módulo de auditoría"))
if rf_anchor and not find_para(doc, lambda p: "agendamiento de citas médicas" in p.text):
    add_after(rf_anchor, [
        ("normal", "El sistema debe permitir el agendamiento de citas médicas por parte del paciente, seleccionando un médico validado, un horario disponible y un motivo de consulta. La cita debe registrarse inicialmente con estado pendiente y debe quedar asociada tanto al paciente como al médico seleccionado (Cayirli & Veral, 2003; Zhao et al., 2017)."),
        ("normal", "El sistema debe permitir que el paciente consulte sus citas registradas y pueda cancelar aquellas que aún no se encuentren completadas o previamente canceladas. La cancelación debe conservar, cuando aplique, un motivo breve que permita documentar la razón de la modificación."),
        ("normal", "El sistema debe permitir que el médico consulte su agenda, filtre citas por fecha, confirme citas pendientes o reprogramadas, reprograme citas a un nuevo horario disponible y marque una cita como completada una vez finalizada la atención."),
        ("normal", "El sistema debe permitir que el médico administre su disponibilidad semanal mediante días, hora de inicio, hora de fin y estado activo, evitando que los pacientes seleccionen horarios fuera de la disponibilidad configurada."),
        ("normal", "El sistema debe permitir que el administrador supervise las citas registradas y filtre por estado, con el fin de contar con una vista general de la operación del sistema y apoyar el seguimiento administrativo."),
    ])

# Add appointment non-functional paragraph.
nf_anchor = find_para(doc, lambda p: p.text.startswith("En términos de disponibilidad, el sistema debe estar accesible"))
if nf_anchor and "gestión de citas" not in nf_anchor.text:
    set_text(nf_anchor, nf_anchor.text + " En el módulo de citas, la disponibilidad también implica que los horarios mostrados al paciente correspondan con espacios vigentes, no ocupados y dentro de la agenda configurada por el médico.")

# Update scope and implementation text to include appointments.
scope_anchor = find_para(doc, lambda p: p.text.startswith("Dentro de este alcance, el sistema permite registrar usuarios"))
if scope_anchor and "gestionar citas" not in scope_anchor.text:
    set_text(scope_anchor, "Dentro de este alcance, el sistema permite registrar usuarios, administrar perfiles, gestionar historiales médicos electrónicos, programar citas médicas y controlar el acceso a la información clínica mediante la generación de tokens de autorización. Estos tokens permiten compartir temporalmente el acceso a la información médica con profesionales de la salud, mientras que el módulo de citas permite ordenar la interacción entre paciente y médico a partir de disponibilidad, horarios y estados de atención.")

impl_anchor = find_para(doc, lambda p: p.text.startswith("El módulo de pacientes permite consultar"))
if impl_anchor and "módulo de citas" not in impl_anchor.text:
    add_after(impl_anchor, [
        ("normal", "El módulo de citas permite al paciente consultar médicos disponibles, seleccionar horarios libres y crear una cita indicando el motivo de consulta. Cada cita se almacena con fecha y hora de inicio, fecha y hora de fin, médico, paciente, motivo, notas opcionales y estado operativo. El flujo inicia con una cita pendiente y posteriormente puede avanzar hacia confirmada, reprogramada, cancelada o completada, según las acciones realizadas por paciente o médico."),
        ("normal", "El módulo de disponibilidad médica permite al médico configurar bloques semanales de atención mediante día de la semana, hora de inicio, hora de fin y estado activo. Con esta información, el sistema genera espacios de atención de treinta minutos y excluye aquellos horarios que ya tienen citas activas. Esta validación evita traslapes y permite que el agendamiento se realice únicamente sobre horarios realmente disponibles (Gupta & Denton, 2008)."),
    ])

# Cases of use narrative: mention appointment cases explicitly.
cases_anchor = find_para(doc, lambda p: p.text.startswith("En conjunto, los casos de uso permiten describir"))
if cases_anchor and "gestión de citas" not in cases_anchor.text:
    set_text(cases_anchor, cases_anchor.text + " Dado que el sistema también incorpora gestión de citas médicas, se agregan casos de uso específicos para disponibilidad, agendamiento, consulta, cancelación, confirmación, reprogramación, cierre y supervisión administrativa de citas.")

# Insert appointment use case tables after CU-09.
template = find_table(doc, "CU-09 – Administrar usuarios y validar médicos") or doc.tables[-1]
existing_titles = {table_title(t) for t in doc.tables}
appointment_cases = [
    [
        ("No. y nombre de CU:", "CU-10 – Gestionar disponibilidad médica"),
        ("Fuentes:", "Requerimientos de agenda médica, control de disponibilidad y prevención de horarios inválidos."),
        ("Actor(es):", "Médico"),
        ("Descripción:", "Permite al médico configurar los días y horarios en los que estará disponible para recibir citas. Cada bloque de disponibilidad incluye día de la semana, hora de inicio, hora de fin y estado activo."),
        ("Pre-Condiciones:", "El médico debe estar autenticado. / El médico debe contar con perfil válido dentro del sistema."),
        ("Flujo básico (escenario de éxito):", "El médico accede al módulo de disponibilidad. / El sistema muestra los horarios registrados. / El médico agrega, modifica o desactiva bloques de atención. / El sistema valida que la hora de inicio sea menor que la hora de fin. / El sistema guarda la disponibilidad actualizada."),
        ("Flujo alternativo (escenarios de fracaso):", "A1. Horario inválido / El sistema solicita corrección. / A2. Usuario sin perfil médico / El sistema deniega la operación. / A3. Error de almacenamiento / El sistema informa la falla."),
        ("Post-Condiciones:", "La disponibilidad semanal del médico queda actualizada y puede ser consultada al momento de agendar citas."),
        ("Notas:", "La disponibilidad médica alimenta la generación de espacios disponibles para pacientes y evita que se programen citas fuera de horario."),
    ],
    [
        ("No. y nombre de CU:", "CU-11 – Agendar cita médica"),
        ("Fuentes:", "Requerimientos de gestión de citas, acceso oportuno a atención médica y disponibilidad de profesionales."),
        ("Actor(es):", "Paciente"),
        ("Descripción:", "Permite al paciente crear una cita médica seleccionando un médico validado, una fecha, un horario disponible, un motivo de consulta y notas opcionales."),
        ("Pre-Condiciones:", "El paciente debe estar autenticado. / Debe existir al menos un médico validado y activo. / El médico seleccionado debe tener disponibilidad en la fecha solicitada."),
        ("Flujo básico (escenario de éxito):", "El paciente accede al módulo de agendar cita. / El sistema muestra médicos disponibles. / El paciente selecciona médico y fecha. / El sistema consulta horarios libres. / El paciente selecciona un horario e ingresa motivo de consulta. / El sistema valida que el horario no esté ocupado. / El sistema registra la cita con estado pendiente."),
        ("Flujo alternativo (escenarios de fracaso):", "A1. No hay horarios disponibles / El sistema muestra mensaje informativo. / A2. Horario ocupado durante la operación / El sistema rechaza la cita y solicita elegir otro horario. / A3. Fecha pasada o inválida / El sistema bloquea el registro."),
        ("Post-Condiciones:", "La cita queda registrada en la agenda del paciente y del médico con estado pendiente. / Se registra el evento en auditoría."),
        ("Notas:", "Cada cita tiene duración fija de treinta minutos en la implementación actual."),
    ],
    [
        ("No. y nombre de CU:", "CU-12 – Consultar y cancelar cita del paciente"),
        ("Fuentes:", "Requerimientos de seguimiento de citas y autonomía del paciente para cancelar solicitudes activas."),
        ("Actor(es):", "Paciente"),
        ("Descripción:", "Permite al paciente visualizar sus citas registradas y cancelar aquellas que aún se encuentren en un estado modificable."),
        ("Pre-Condiciones:", "El paciente debe estar autenticado. / Debe existir al menos una cita asociada al paciente."),
        ("Flujo básico (escenario de éxito):", "El paciente accede al módulo Mis Citas. / El sistema muestra las citas ordenadas por fecha. / El paciente selecciona una cita activa. / El paciente confirma la cancelación e ingresa un motivo opcional. / El sistema cambia el estado de la cita a cancelada. / El sistema registra la acción en auditoría."),
        ("Flujo alternativo (escenarios de fracaso):", "A1. Cita ya completada o cancelada / El sistema informa que ya no puede cancelarse. / A2. Cita perteneciente a otro paciente / El sistema deniega la operación. / A3. Error de actualización / El sistema muestra mensaje de error."),
        ("Post-Condiciones:", "La cita queda en estado cancelada y ya no se considera como horario ocupado para nuevas solicitudes."),
        ("Notas:", "La cancelación no elimina el registro; conserva evidencia del evento y del motivo cuando se proporciona."),
    ],
    [
        ("No. y nombre de CU:", "CU-13 – Gestionar agenda médica"),
        ("Fuentes:", "Requerimientos de operación del médico, seguimiento de citas asignadas y continuidad de atención."),
        ("Actor(es):", "Médico"),
        ("Descripción:", "Permite al médico consultar su agenda, filtrar citas por fecha y ejecutar acciones sobre citas asignadas, como confirmar, reprogramar o marcar como completadas."),
        ("Pre-Condiciones:", "El médico debe estar autenticado. / El médico debe tener citas asociadas a su perfil o el sistema debe mostrar un estado vacío."),
        ("Flujo básico (escenario de éxito):", "El médico accede al módulo Agenda. / El sistema muestra citas asociadas al médico. / El médico filtra por fecha si lo requiere. / El médico selecciona una cita pendiente o reprogramada. / El médico confirma la cita, la reprograma a un horario válido o la marca como completada. / El sistema actualiza el estado y registra el evento."),
        ("Flujo alternativo (escenarios de fracaso):", "A1. Cita cancelada o completada / El sistema bloquea acciones no permitidas. / A2. Nuevo horario fuera de disponibilidad / El sistema rechaza la reprogramación. / A3. Cita no asignada al médico / El sistema deniega el acceso."),
        ("Post-Condiciones:", "La cita queda actualizada con el estado correspondiente: confirmada, reprogramada o completada."),
        ("Notas:", "La agenda médica se relaciona con la disponibilidad configurada por el médico y con las citas creadas por los pacientes."),
    ],
    [
        ("No. y nombre de CU:", "CU-14 – Supervisar citas desde administración"),
        ("Fuentes:", "Requerimientos de supervisión administrativa y seguimiento operativo del sistema."),
        ("Actor(es):", "Administrador"),
        ("Descripción:", "Permite al administrador consultar las citas registradas en el sistema y filtrarlas por estado para obtener una visión general de la operación."),
        ("Pre-Condiciones:", "El administrador debe estar autenticado. / Deben existir citas registradas o el sistema debe mostrar un estado informativo vacío."),
        ("Flujo básico (escenario de éxito):", "El administrador accede al módulo de citas. / El sistema muestra citas con información de paciente, médico, fecha, motivo y estado. / El administrador aplica filtros por estado si es necesario. / El sistema actualiza la tabla con los resultados correspondientes."),
        ("Flujo alternativo (escenarios de fracaso):", "A1. Estado de filtro inválido / El sistema rechaza la consulta. / A2. Usuario sin rol administrador / El sistema deniega el acceso. / A3. No existen citas / El sistema muestra mensaje informativo."),
        ("Post-Condiciones:", "El administrador obtiene una vista de seguimiento sin modificar directamente la información clínica ni el historial médico."),
        ("Notas:", "La supervisión administrativa ayuda a observar carga de citas pendientes, confirmadas y completadas dentro del prototipo."),
    ],
]

last_element = template._element
for rows in appointment_cases:
    title = rows[0][1]
    if title in existing_titles:
        continue
    spacer = OxmlElement("w:p")
    last_element.addnext(spacer)
    last_element = spacer
    new_table = insert_table_after(last_element, template, rows)
    last_element = new_table._element
    existing_titles.add(title)

# Add/update bibliography references, sorted.
bib = find_para(doc, lambda p: p.text.strip().startswith("Bibliograf"))
anexos = find_para(doc, lambda p: p.text.strip() == "Anexos")
if bib and anexos:
    refs = remove_bibliography_between(doc, bib, anexos)
    merged = sorted(set(refs + APPOINTMENT_REFS), key=lambda x: x.lower())
    for ref in merged:
        para = paragraph_before(anexos, ref, "normal")
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.first_line_indent = Inches(-0.5)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.0
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(10)

doc.save(DOCX)
print(f"Updated appointments content in {DOCX}")
