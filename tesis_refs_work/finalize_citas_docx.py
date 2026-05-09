from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


DOCX = Path(r"C:\Users\julio\Gestor_CHM\MedRecord\tesis_refs_work\Plantilla_de_tesis_citas_actualizada.docx")


doc = Document(DOCX)

anexos = None
for p in doc.paragraphs:
    if p.text.strip() == "Anexos":
        anexos = p
        break

if anexos is None:
    raise SystemExit("No se encontró Anexos")

prev = anexos._element.getprevious()
has_page_break = False
if prev is not None:
    prev_para = Paragraph(prev, anexos._parent)
    has_page_break = bool(prev_para._element.xpath(".//w:br[@w:type='page']"))

if not has_page_break:
    p_el = OxmlElement("w:p")
    anexos._element.addprevious(p_el)
    paragraph = Paragraph(p_el, anexos._parent)
    paragraph.add_run().add_break(WD_BREAK.PAGE)

doc.save(DOCX)
print(f"Finalized {DOCX}")
