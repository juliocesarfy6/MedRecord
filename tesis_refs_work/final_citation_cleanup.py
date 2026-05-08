from pathlib import Path

from docx import Document
from docx.shared import Pt


DOCX = Path(r"C:\Users\julio\Gestor_CHM\MedRecord\tesis_refs_work\Plantilla_de_tesis_con_referencias.docx")


doc = Document(DOCX)

replacements = {
    "(World Health Organization, 2021) (Tang et al., 2006; Poss-Doering et al., 2018; Baudendistel et al., 2015)": "(World Health Organization, 2021; Tang et al., 2006; Poss-Doering et al., 2018; Baudendistel et al., 2015)",
}

for paragraph in doc.paragraphs:
    text = paragraph.text
    new_text = text
    for old, new in replacements.items():
        new_text = new_text.replace(old, new)
    if new_text != text:
        style = paragraph.style
        alignment = paragraph.alignment
        for run in list(paragraph.runs):
            run._element.getparent().remove(run._element)
        run = paragraph.add_run(new_text)
        run.font.name = "Arial"
        run.font.size = Pt(12)
        paragraph.style = style
        paragraph.alignment = alignment

doc.save(DOCX)
print(f"Cleaned {DOCX}")
