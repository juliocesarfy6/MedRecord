from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Pt


DOCX = Path(r"C:\Users\julio\Gestor_CHM\MedRecord\tesis_refs_work\Plantilla_de_tesis_con_referencias.docx")


REFERENCES = [
    "Adler-Milstein, J., & Jha, A. K. (2017). HITECH Act drove large gains in hospital electronic health record adoption. Health Affairs, 36(8), 1416-1422. https://doi.org/10.1377/hlthaff.2016.1651",
    "Bass, L., Clements, P., & Kazman, R. (2021). Software architecture in practice (4th ed.). Addison-Wesley.",
    "Bates, D. W., Leape, L. L., Cullen, D. J., Laird, N., Petersen, L. A., Teich, J. M., Burdick, E., Hickey, M., Kleefield, S., Shea, B., Vander Vliet, M., & Seger, D. L. (1998). Effect of computerized physician order entry and a team intervention on prevention of serious medication errors. JAMA, 280(15), 1311-1316. https://doi.org/10.1001/jama.280.15.1311",
    "Baudendistel, I., Winkler, E. C., Kamradt, M., Brophy, S., Längst, G., Eckrich, F., Heinze, O., Bergh, B., Szecsenyi, J., & Ose, D. (2015). The patients' active role in managing a personal electronic health record: A qualitative analysis. Supportive Care in Cancer, 23(9), 2613-2621. https://doi.org/10.1007/s00520-015-2620-1",
    "Benson, T., & Grieve, G. (2016). Principles of health interoperability: SNOMED CT, HL7 and FHIR (3rd ed.). Springer.",
    "Blumenthal, D., & Tavenner, M. (2010). The meaningful use regulation for electronic health records. New England Journal of Medicine, 363(6), 501-504. https://doi.org/10.1056/NEJMp1006114",
    "Boonstra, A., & Broekhuis, M. (2010). Barriers to the acceptance of electronic medical records by physicians: From systematic review to taxonomy and interventions. BMC Health Services Research, 10, Article 231. https://doi.org/10.1186/1472-6963-10-231",
    "Chacon, S., & Straub, B. (2014). Pro Git (2nd ed.). Apress.",
    "Cimino, J. J. (2013). Improving the electronic health record: Are clinicians getting what they wished for? JAMA, 309(10), 991-992. https://doi.org/10.1001/jama.2013.890",
    "Diario Oficial de la Federación. (2010, 5 de julio). Ley Federal de Protección de Datos Personales en Posesión de los Particulares. https://www.diputados.gob.mx/LeyesBiblio/pdf/LFPDPPP.pdf",
    "Elmasri, R., & Navathe, S. B. (2016). Fundamentals of database systems (7th ed.). Pearson.",
    "Ferraiolo, D. F., Sandhu, R., Gavrila, S., Kuhn, D. R., & Chandramouli, R. (2001). Proposed NIST standard for role-based access control. ACM Transactions on Information and System Security, 4(3), 224-274. https://doi.org/10.1145/501978.501980",
    "Fielding, R. T. (2000). Architectural styles and the design of network-based software architectures [Doctoral dissertation, University of California, Irvine]. https://www.ics.uci.edu/~fielding/pubs/dissertation/top.htm",
    "Fowler, M. (2002). Patterns of enterprise application architecture. Addison-Wesley.",
    "Gamma, E., Helm, R., Johnson, R., & Vlissides, J. (1994). Design patterns: Elements of reusable object-oriented software. Addison-Wesley.",
    "Giardina, T. D., Menon, S., Parrish, D. E., Sittig, D. F., & Singh, H. (2014). Patient access to medical records and healthcare outcomes: A systematic review. Journal of the American Medical Informatics Association, 21(4), 737-741. https://doi.org/10.1136/amiajnl-2013-002239",
    "GitHub. (2024). GitHub Docs: About version control and Git. https://docs.github.com/",
    "Goldzweig, C. L., Towfigh, A., Maglione, M., & Shekelle, P. G. (2009). Costs and benefits of health information technology: New trends from the literature. Health Affairs, 28(2), w282-w293. https://doi.org/10.1377/hlthaff.28.2.w282",
    "Hardt, D. (2012). The OAuth 2.0 authorization framework (RFC 6749). Internet Engineering Task Force. https://doi.org/10.17487/RFC6749",
    "HL7 International. (2019). FHIR Release 4: Fast Healthcare Interoperability Resources. https://hl7.org/fhir/R4/",
    "IEEE Computer Society. (2014). Guide to the Software Engineering Body of Knowledge (SWEBOK Guide), Version 3.0. IEEE Computer Society.",
    "ISO/IEC. (2022a). ISO/IEC 27001:2022: Information security, cybersecurity and privacy protection: Information security management systems: Requirements. International Organization for Standardization.",
    "ISO/IEC. (2022b). ISO/IEC 27002:2022: Information security, cybersecurity and privacy protection: Information security controls. International Organization for Standardization.",
    "Jha, A. K., DesRoches, C. M., Campbell, E. G., Donelan, K., Rao, S. R., Ferris, T. G., Shields, A., Rosenbaum, S., & Blumenthal, D. (2009). Use of electronic health records in U.S. hospitals. New England Journal of Medicine, 360(16), 1628-1638. https://doi.org/10.1056/NEJMsa0900592",
    "Jones, M., Bradley, J., & Sakimura, N. (2015). JSON Web Token (JWT) (RFC 7519). Internet Engineering Task Force. https://doi.org/10.17487/RFC7519",
    "Kaushal, R., & Bates, D. W. (2003). Effects of computerized physician order entry and clinical decision support systems on medication safety: A systematic review. Archives of Internal Medicine, 163(12), 1409-1416. https://doi.org/10.1001/archinte.163.12.1409",
    "Kripalani, S., LeFevre, F., Phillips, C. O., Williams, M. V., Basaviah, P., & Baker, D. W. (2007). Deficits in communication and information transfer between hospital-based and primary care physicians: Implications for patient safety and continuity of care. JAMA, 297(8), 831-841. https://doi.org/10.1001/jama.297.8.831",
    "Kruse, C. S., Stein, A., Thomas, H., & Kaur, H. (2018). The use of electronic health records to support population health: A systematic review of the literature. Journal of Medical Systems, 42, Article 214. https://doi.org/10.1007/s10916-018-1075-6",
    "Menachemi, N., & Collum, T. H. (2011). Benefits and drawbacks of electronic health record systems. Risk Management and Healthcare Policy, 4, 47-55. https://doi.org/10.2147/RMHP.S12985",
    "Mozilla Developer Network. (2024). Web technology for developers. https://developer.mozilla.org/",
    "National Institute of Standards and Technology. (2020). Security and privacy controls for information systems and organizations (SP 800-53 Rev. 5). https://doi.org/10.6028/NIST.SP.800-53r5",
    "National Institute of Standards and Technology. (2024). The NIST Cybersecurity Framework (CSF) 2.0. https://doi.org/10.6028/NIST.CSWP.29",
    "Nielsen, J. (1994). Usability engineering. Morgan Kaufmann.",
    "Norman, D. A. (2013). The design of everyday things (Revised and expanded ed.). Basic Books.",
    "Object Management Group. (2017). OMG Unified Modeling Language (OMG UML), Version 2.5.1. https://www.omg.org/spec/UML/2.5.1/",
    "Open Source Initiative. (2007). The open source definition. https://opensource.org/osd/",
    "OWASP Foundation. (2021). OWASP Top 10: The ten most critical web application security risks. https://owasp.org/Top10/",
    "Pagliari, C., Detmer, D., & Singleton, P. (2007). Potential of electronic personal health records. BMJ, 335(7615), 330-333. https://doi.org/10.1136/bmj.39279.482963.AD",
    "Poss-Doering, R., Kunz, A., Pohlmann, S., Hofmann, H., Kiel, M., Winkler, E. C., Ose, D., & Szecsenyi, J. (2018). Utilizing a prototype patient-controlled electronic health record in Germany: Qualitative analysis of user-reported perceptions and perspectives. JMIR Formative Research, 2(2), e10411. https://doi.org/10.2196/10411",
    "PostgreSQL Global Development Group. (2024). PostgreSQL documentation. https://www.postgresql.org/docs/",
    "Pressman, R. S., & Maxim, B. R. (2020). Software engineering: A practitioner's approach (9th ed.). McGraw-Hill Education.",
    "Saltzer, J. H., & Schroeder, M. D. (1975). The protection of information in computer systems. Proceedings of the IEEE, 63(9), 1278-1308. https://doi.org/10.1109/PROC.1975.9939",
    "Schwaber, K., & Sutherland, J. (2020). The Scrum Guide: The definitive guide to Scrum. https://scrumguides.org/scrum-guide.html",
    "Secretaría de Salud. (2012a). Norma Oficial Mexicana NOM-024-SSA3-2012, Sistemas de información de registro electrónico para la salud. Diario Oficial de la Federación.",
    "Secretaría de Salud. (2012b). Norma Oficial Mexicana NOM-004-SSA3-2012, Del expediente clínico. Diario Oficial de la Federación.",
    "Silberschatz, A., Korth, H. F., & Sudarshan, S. (2020). Database system concepts (7th ed.). McGraw-Hill Education.",
    "Sittig, D. F., & Singh, H. (2012). Rights and responsibilities of users of electronic health records. CMAJ, 184(13), 1479-1483. https://doi.org/10.1503/cmaj.111599",
    "Sommerville, I. (2016). Software engineering (10th ed.). Pearson.",
    "Tang, P. C., Ash, J. S., Bates, D. W., Overhage, J. M., & Sands, D. Z. (2006). Personal health records: Definitions, benefits, and strategies for overcoming barriers to adoption. Journal of the American Medical Informatics Association, 13(2), 121-126. https://doi.org/10.1197/jamia.M2025",
    "Vest, J. R., & Gamm, L. D. (2010). Health information exchange: Persistent challenges and new strategies. Journal of the American Medical Informatics Association, 17(3), 288-294. https://doi.org/10.1136/jamia.2010.003673",
    "World Health Organization. (2021). Global strategy on digital health 2020-2025. World Health Organization. https://www.who.int/publications/i/item/9789240020924",
]


REMOVE_TEXTS = {
    "Referencia",
    "Bates, D. W., et al. (2003).\nEffect of computerized physician order entry and a team intervention on prevention of serious medication errors.\nJAMA.",
    "Bates, D. W., et al. (2003).",
    "Effect of computerized physician order entry and a team intervention on prevention of serious medication errors.",
    "JAMA.",
    "HL7 International. (2019).\nFHIR Release 4: Fast Healthcare Interoperability Resources.",
    "World Health Organization (WHO). (2016).\nElectronic Health Records: Manual for Developing Countries.",
    "World Health Organization (WHO). (2016).",
    "Electronic Health Records: Manual for Developing Countries.",
    "ISO/IEC 27001. (2013).",
    "Information Security Management Systems Requirements.",
}


CITATIONS = [
    ("La adopción de Expedientes Clínicos Electrónicos", "(Bates et al., 1998; Menachemi & Collum, 2011; Kruse et al., 2018)"),
    ("Cada hospital, clínica o consultorio mantiene", "(Vest & Gamm, 2010; Adler-Milstein & Jha, 2017; HL7 International, 2019)"),
    ("la fragmentación de datos clínicos constituye", "(Kripalani et al., 2007; Sittig & Singh, 2012)"),
    ("no existe un modelo ampliamente adoptado", "(Secretaría de Salud, 2012a; Diario Oficial de la Federación, 2010)"),
    ("Falta de acceso al historial clínico completo", "(Cimino, 2013; Giardina et al., 2014)"),
    ("estándares de interoperabilidad como HL7 y FHIR", "(HL7 International, 2019; Benson & Grieve, 2016)"),
    ("dato sensible protegido por la Ley Federal", "(Diario Oficial de la Federación, 2010; Secretaría de Salud, 2012b)"),
    ("enfoque centrado en el paciente", "(Tang et al., 2006; Poss-Doering et al., 2018; Baudendistel et al., 2015)"),
    ("gestión interna de hospitales o clínicas", "(Boonstra & Broekhuis, 2010; Goldzweig et al., 2009)"),
    ("confidencialidad, integridad y disponibilidad", "(ISO/IEC, 2022a; NIST, 2020)"),
    ("control de acceso, registro de actividad", "(NIST, 2020; OWASP Foundation, 2021; Saltzer & Schroeder, 1975)"),
    ("herramientas, tecnologías y marcos de desarrollo", "(Sommerville, 2016; Pressman & Maxim, 2020)"),
    ("arquitectura en capas", "(Bass et al., 2021; Fowler, 2002)"),
    ("gestor de base de datos", "(Silberschatz et al., 2020; Elmasri & Navathe, 2016)"),
    ("autorizaciones temporales mediante tokens", "(Hardt, 2012; Jones et al., 2015; Ferraiolo et al., 2001)"),
    ("aplicación web", "(Fielding, 2000; Mozilla Developer Network, 2024)"),
    ("tecnologías de código abierto", "(Open Source Initiative, 2007; PostgreSQL Global Development Group, 2024)"),
    ("arquitectura modular basada en capas", "(Bass et al., 2021; Gamma et al., 1994)"),
    ("metodología de trabajo orientada al desarrollo iterativo", "(Schwaber & Sutherland, 2020; Sommerville, 2016)"),
    ("diagramas UML", "(Object Management Group, 2017)"),
    ("control de versiones mediante la plataforma GitHub", "(Chacon & Straub, 2014; GitHub, 2024)"),
    ("SDLC", "(Pressman & Maxim, 2020; IEEE Computer Society, 2014)"),
    ("interfaz clara e intuitiva", "(Nielsen, 1994; Norman, 2013)"),
    ("módulo de auditoría", "(NIST, 2020; ISO/IEC, 2022b)"),
    ("datos médicos sensibles", "(Diario Oficial de la Federación, 2010; ISO/IEC, 2022a)"),
]


def remove_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def set_paragraph_text(paragraph, text):
    style = paragraph.style
    alignment = paragraph.alignment
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    paragraph.style = style
    paragraph.alignment = alignment


def add_citation(paragraph, citation):
    if citation in paragraph.text:
        return
    text = paragraph.text.strip()
    if text.endswith("."):
        text = f"{text[:-1]} {citation}."
    else:
        text = f"{text} {citation}"
    set_paragraph_text(paragraph, text)


def insert_paragraph_before(target, text, style="normal"):
    new_p = OxmlElement("w:p")
    target._element.addprevious(new_p)
    paragraph = Paragraph(new_p, target._parent)
    paragraph.style = style
    paragraph.add_run(text)
    return paragraph


doc = Document(DOCX)

TEXT_REPLACEMENTS = {
    "(WHO, 2016)": "(World Health Organization, 2021)",
    "(ISO/IEC 27001, 2013)": "(ISO/IEC, 2022a)",
    "(Bates et al., 2003)": "(Bates et al., 1998; Menachemi & Collum, 2011; Kruse et al., 2018)",
}

for paragraph in list(doc.paragraphs):
    if paragraph.text.strip() in REMOVE_TEXTS:
        remove_paragraph(paragraph)

for paragraph in doc.paragraphs:
    original = paragraph.text
    updated = original
    for old, new in TEXT_REPLACEMENTS.items():
        updated = updated.replace(old, new)
    if updated != original:
        set_paragraph_text(paragraph, updated)

for needle, citation in CITATIONS:
    for paragraph in doc.paragraphs:
        if needle in paragraph.text and citation not in paragraph.text:
            add_citation(paragraph, citation)
            break

# Remove any old empty bibliography placeholder content, then insert the APA-style list before Anexos.
bib = None
anexos = None
for paragraph in doc.paragraphs:
    normalized = paragraph.text.strip().lower()
    if normalized == "bibliografía":
        bib = paragraph
    if normalized == "anexos":
        anexos = paragraph
        break

if bib is None:
    if anexos is None:
        bib = doc.add_paragraph("Bibliografía", style="Heading 2")
    else:
        bib = insert_paragraph_before(anexos, "Bibliografía", style="Heading 2")
else:
    bib.style = "Heading 2"
    bib.alignment = WD_ALIGN_PARAGRAPH.LEFT

if anexos is None:
    anexos = doc.add_paragraph("Anexos", style="Heading 2")

between = []
collecting = False
for paragraph in list(doc.paragraphs):
    if paragraph._element is bib._element:
        collecting = True
        continue
    if paragraph._element is anexos._element:
        break
    if collecting:
        between.append(paragraph)
for paragraph in between:
    remove_paragraph(paragraph)

for ref in REFERENCES:
    paragraph = insert_paragraph_before(anexos, ref, style="normal")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.5)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.0
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10)

doc.save(DOCX)
print(f"Saved {DOCX}")
print(f"References: {len(REFERENCES)}")
