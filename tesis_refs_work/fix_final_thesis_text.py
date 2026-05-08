from pathlib import Path

from docx import Document
from docx.shared import Pt


DOCX = Path(r"C:\Users\julio\Gestor_CHM\MedRecord\tesis_refs_work\Plantilla_de_tesis_ampliada.docx")


doc = Document(DOCX)

replacements = {
    "Actualmente, el prototipo implementado incluye únicamente el módulo de autenticación. En etapas posteriores del desarrollo se integrarán las interfaces correspondientes a la gestión del historial médico, autorización explícita de acceso a médicos y registro de trazabilidad de consultas, manteniendo los mismos principios de diseño, validación y claridad visual aplicados en estas pantallas iniciales.": "El prototipo implementado ya contempla módulos base para autenticación, perfil del paciente, historial médico, generación de tokens y auditoría. Las interfaces muestran una organización orientada a que el paciente pueda identificar rápidamente sus consultas registradas, tokens activos y eventos recientes de acceso. Como trabajo complementario dentro del mismo alcance, se propone completar la revocación de tokens desde el backend, reforzar la validación del nivel de acceso antes de permitir edición de registros y ampliar las vistas del médico y del administrador para cubrir todos los casos de uso definidos.",
}

for paragraph in doc.paragraphs:
    text = paragraph.text
    if text in replacements:
        style = paragraph.style
        alignment = paragraph.alignment
        for run in list(paragraph.runs):
            run._element.getparent().remove(run._element)
        run = paragraph.add_run(replacements[text])
        run.font.name = "Arial"
        run.font.size = Pt(12)
        paragraph.style = style
        paragraph.alignment = alignment

doc.save(DOCX)
print(f"Fixed final text in {DOCX}")
