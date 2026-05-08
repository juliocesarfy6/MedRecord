from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt


DOCX = Path(r"C:\Users\julio\Gestor_CHM\MedRecord\tesis_refs_work\Plantilla_de_tesis_con_referencias.docx")


def remove_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


doc = Document(DOCX)

bib = None
anexos = None
for paragraph in doc.paragraphs:
    text = paragraph.text.strip().lower()
    if text == "bibliografía":
        bib = paragraph
    if text == "anexos":
        anexos = paragraph
        break

if bib and anexos:
    collecting = False
    for paragraph in doc.paragraphs:
        if paragraph._element is bib._element:
            collecting = True
            continue
        if paragraph._element is anexos._element:
            break
        if collecting and paragraph.text.strip():
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.first_line_indent = Inches(-0.5)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = Pt(10)

    # Force Anexos to start on its own page after the bibliography.
    previous = anexos._element.getprevious()
    if previous is not None:
        page_break_paragraph = doc.add_paragraph()
        page_break_paragraph.add_run().add_break(WD_BREAK.PAGE)
        anexos._element.addprevious(page_break_paragraph._element)

# Remove empty paragraphs at the very end of the document.
for paragraph in reversed(doc.paragraphs):
    if paragraph.text.strip():
        break
    remove_paragraph(paragraph)

doc.save(DOCX)
print(f"Polished {DOCX}")
