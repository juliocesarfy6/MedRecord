from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


DOCX = Path(r"C:\Users\julio\Gestor_CHM\MedRecord\tesis_refs_work\Plantilla_de_tesis_ampliada.docx")


def paragraph_after(paragraph, text="", style="normal"):
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    p.style = style
    if text:
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(12)
    return p


def set_text(paragraph, text, style=None):
    if style:
        paragraph.style = style
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)


def remove_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def find_para(doc, predicate):
    for p in doc.paragraphs:
        if predicate(p):
            return p
    return None


def insert_block_after(anchor, items):
    current = anchor
    for style, text in reversed([]):
        pass
    for style, text in items:
        current = paragraph_after(current, text, style)
        if style.startswith("Heading"):
            for run in current.runs:
                run.bold = True
                run.font.size = Pt(12 if style == "Heading 3" else 14)
        current.paragraph_format.space_after = Pt(6)
    return current


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def format_table(table):
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(9)
            if row_idx == 0:
                shade_cell(cell, "EAF2F8")
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True


def set_table_field(table, field_name, value):
    for row in table.rows:
        if row.cells[0].text.strip().startswith(field_name):
            row.cells[1].text = value
            return


def append_use_case_table(doc, anchor_paragraph, rows):
    table = doc.add_table(rows=0, cols=2)
    for label, value in rows:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
    format_table(table)
    anchor_paragraph._element.addnext(table._element)
    caption = paragraph_after(anchor_paragraph, "")
    return table


doc = Document(DOCX)

# Resumen: replace the placeholder dot with a real abstract.
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "Resumen":
        next_p = doc.paragraphs[i + 1]
        if next_p.text.strip() == ".":
            set_text(
                next_p,
                "El presente proyecto propone el diseño e implementación de un sistema web para la gestión del historial médico electrónico desde un enfoque centrado en el paciente. La problemática principal que atiende la propuesta es la fragmentación de la información clínica, situación que se presenta cuando los datos médicos de una persona permanecen distribuidos entre consultorios, instituciones o especialistas sin un mecanismo práctico que permita integrarlos y compartirlos de forma controlada. Esta dispersión limita la continuidad de la atención, dificulta la consulta de antecedentes relevantes y puede incrementar riesgos asociados a duplicidad de estudios, decisiones clínicas con información incompleta y falta de trazabilidad sobre el uso de los datos.",
            )
            paragraph_after(
                next_p,
                "Para responder a esta problemática, el sistema plantea una arquitectura web compuesta por una capa de presentación, una capa de lógica de negocio y una capa de datos. El prototipo contempla usuarios con roles diferenciados, principalmente paciente, médico y administrador. El paciente puede consultar y actualizar su información básica, visualizar su historial médico, generar tokens temporales de acceso y revisar registros de auditoría. El médico puede acceder al expediente únicamente cuando cuenta con una autorización válida y, según el nivel de permiso concedido, consultar información clínica o registrar una nueva consulta. El administrador, por su parte, conserva funciones de control relacionadas con la gestión de usuarios y validación de profesionales de la salud.",
            )
            paragraph_after(
                doc.paragraphs[i + 2],
                "El modelo de autorización se apoya en tokens únicos con vigencia temporal y nivel de acceso definido, lo cual permite que el paciente mantenga control explícito sobre quién puede consultar su información. Además, el sistema incorpora autenticación, control de roles y registros de auditoría como mecanismos orientados a fortalecer la confidencialidad, integridad, disponibilidad y trazabilidad de los datos médicos. La propuesta no pretende sustituir sistemas hospitalarios oficiales ni realizar diagnósticos automatizados; su finalidad es demostrar, mediante un prototipo académico funcional, la viabilidad de un modelo de gestión de historial médico seguro, organizado y centrado en el paciente.",
            )
        break

# Marco teorico: replace placeholders with developed subsections before related works.
for p in list(doc.paragraphs):
    if p.text.strip() in {"2.1 Temas relacionados", "2.2 Trabajos relacionados"}:
        remove_paragraph(p)

marco = find_para(doc, lambda p: p.text.strip() == "Marco Teórico")
first_related = find_para(doc, lambda p: p.text.strip().startswith("2.2.1 Poss-Doering"))
if marco and first_related:
    theoretical_items = [
        ("Heading 3", "2.1 Temas relacionados"),
        (
            "normal",
            "El marco teórico del proyecto se organiza a partir de cinco conceptos principales: expediente clínico electrónico, historial o registro personal de salud, interoperabilidad, seguridad de la información y arquitectura de sistemas web. Estos conceptos permiten fundamentar por qué el sistema propuesto no se limita a almacenar datos médicos, sino que busca resolver un problema de acceso, control y trazabilidad. En el contexto de la salud digital, la utilidad de una plataforma no depende únicamente de que exista una base de datos, sino de que la información pueda consultarse de manera oportuna, protegerse frente a accesos no autorizados y mantenerse vinculada al paciente durante su recorrido por distintos servicios de atención (Menachemi & Collum, 2011; Kruse et al., 2018).",
        ),
        ("Heading 3", "2.1.1 Expediente clínico electrónico y fragmentación de información"),
        (
            "normal",
            "El expediente clínico electrónico se entiende como una representación digital de la información médica generada durante la atención de un paciente. Puede contener datos de identificación, antecedentes, diagnósticos, tratamientos, observaciones clínicas, resultados de estudios y registros de consulta. A diferencia del expediente físico, el formato electrónico facilita la consulta, actualización y almacenamiento estructurado de la información; sin embargo, su adopción no garantiza por sí misma continuidad asistencial. Cuando cada institución o profesional utiliza sistemas independientes, los datos pueden quedar aislados y perder valor clínico al no estar disponibles en el momento de la atención (Bates et al., 1998; Cimino, 2013).",
        ),
        (
            "normal",
            "La fragmentación del expediente médico es especialmente relevante en pacientes que reciben atención de varios especialistas o que cambian de institución. En estos casos, el médico puede no conocer tratamientos previos, alergias, diagnósticos anteriores o estudios recientes, lo cual puede derivar en duplicidad de pruebas, retrasos o decisiones basadas en información incompleta. Desde la ingeniería de software, este problema puede analizarse como una falla de disponibilidad, integración y gobierno de datos: la información existe, pero no se encuentra organizada en un punto accesible, ni cuenta con reglas claras para su consulta segura (Kripalani et al., 2007; Vest & Gamm, 2010).",
        ),
        (
            "normal",
            "El sistema propuesto aborda esta problemática mediante una plataforma centralizada para fines académicos, en la que el historial médico del paciente se almacena en una base de datos relacional y se consulta mediante módulos web. Aunque el prototipo no se integra directamente con hospitales reales, su diseño permite representar el flujo básico de un expediente consultable por el paciente y compartible con profesionales de la salud bajo autorización explícita. Esta decisión delimita el alcance del proyecto y al mismo tiempo permite validar los componentes esenciales del modelo: registro, consulta, autorización temporal y trazabilidad.",
        ),
        ("Heading 3", "2.1.2 Registros personales de salud y enfoque centrado en el paciente"),
        (
            "normal",
            "Los registros personales de salud, conocidos como Personal Health Records, proponen que el paciente tenga un papel más activo en la gestión de su información médica. Este enfoque no significa que el paciente sustituya al profesional de la salud, sino que cuente con herramientas para consultar, organizar y compartir información relevante cuando sea necesario. Diversos estudios han señalado que el acceso del paciente a su expediente puede fortalecer la percepción de control, mejorar la comunicación con los profesionales y favorecer una participación más informada en el proceso de atención (Tang et al., 2006; Giardina et al., 2014).",
        ),
        (
            "normal",
            "En esta tesis, el enfoque centrado en el paciente se materializa en dos decisiones de diseño. La primera es que el paciente conserva acceso directo a su perfil, historial médico, tokens generados y auditoría de accesos. La segunda es que el acceso del médico no se asume como permanente, sino como una autorización que debe ser emitida por el propio paciente. Esta lógica se diferencia de modelos institucionales tradicionales, donde el control suele recaer en la organización que resguarda el expediente. El objetivo es que el paciente tenga una participación clara en la decisión de quién consulta su información, por cuánto tiempo y con qué nivel de permiso.",
        ),
        (
            "normal",
            "El modelo no elimina la responsabilidad clínica del médico ni convierte al paciente en administrador técnico del sistema. Más bien, proporciona una interfaz simplificada para ejercer control sobre acciones concretas: generar un token, definir vigencia, seleccionar nivel de acceso, compartir el código con un profesional y revisar posteriormente los eventos asociados. De esta manera, el diseño busca equilibrar autonomía, facilidad de uso y seguridad, evitando que el control del paciente dependa de procesos administrativos complejos.",
        ),
        ("Heading 3", "2.1.3 Interoperabilidad y estándares de información en salud"),
        (
            "normal",
            "La interoperabilidad se refiere a la capacidad de diferentes sistemas para intercambiar, interpretar y utilizar información de manera coherente. En salud, este concepto es fundamental porque los datos clínicos pueden generarse en laboratorios, consultorios, hospitales, farmacias o plataformas de seguimiento. Estándares como HL7 y FHIR buscan facilitar la representación estructurada de recursos clínicos y la comunicación entre sistemas heterogéneos (Benson & Grieve, 2016; HL7 International, 2019).",
        ),
        (
            "normal",
            "El prototipo desarrollado no implementa de forma completa HL7 FHIR, pero toma el estándar como referencia conceptual para organizar la información alrededor de entidades clínicas identificables, tales como paciente, médico, registro médico y eventos de acceso. Esta decisión es importante porque permite que el sistema no sea concebido como una aplicación aislada sin proyección futura, sino como una base que podría evolucionar hacia esquemas de intercambio más formales. En una versión posterior, los registros médicos podrían mapearse a recursos FHIR y exponerse mediante servicios compatibles con integraciones externas.",
        ),
        ("Heading 3", "2.1.4 Seguridad, control de acceso y trazabilidad"),
        (
            "normal",
            "La información médica pertenece a una categoría especialmente sensible, ya que puede revelar condiciones de salud, tratamientos, diagnósticos, antecedentes personales y otros datos que afectan directamente la privacidad del individuo. Por ello, el diseño de un sistema de historial médico debe considerar principios de confidencialidad, integridad y disponibilidad. La confidencialidad implica restringir el acceso a usuarios autorizados; la integridad exige preservar la exactitud de los datos; y la disponibilidad busca que la información pueda consultarse cuando se requiere para apoyar la atención (ISO/IEC, 2022a; NIST, 2020).",
        ),
        (
            "normal",
            "El control de acceso implementado en el prototipo combina autenticación, autorización basada en roles y tokens temporales. La autenticación permite identificar al usuario mediante credenciales y sesión segura; la autorización basada en roles diferencia las acciones disponibles para pacientes, médicos y administradores; y el token temporal funciona como un mecanismo de permiso explícito emitido por el paciente. Este diseño se relaciona con modelos de control de acceso que buscan asignar privilegios de acuerdo con funciones y condiciones específicas del sistema (Ferraiolo et al., 2001; Hardt, 2012; Jones et al., 2015).",
        ),
        (
            "normal",
            "La trazabilidad complementa el control de acceso porque no basta con permitir o negar operaciones: también es necesario dejar evidencia de los eventos relevantes. En el sistema, la auditoría registra acciones como consulta de perfil, actualización de datos, creación de registros médicos o accesos al historial. Estos registros permiten responder preguntas como quién realizó una acción, cuándo ocurrió y sobre qué paciente se ejecutó. Para un sistema de salud, esta información es valiosa porque fortalece la transparencia y permite detectar usos indebidos o comportamientos anómalos (NIST, 2020; ISO/IEC, 2022b).",
        ),
        ("Heading 3", "2.1.5 Arquitectura web del sistema"),
        (
            "normal",
            "La arquitectura del sistema se organiza en tres capas: presentación, lógica de negocio y datos. La capa de presentación corresponde a la aplicación web desarrollada con Angular, donde el usuario interactúa con formularios, tablas, paneles y vistas específicas de su rol. La capa de lógica de negocio se implementa mediante NestJS, encargado de procesar solicitudes, aplicar reglas de autorización, validar tokens y coordinar operaciones con los servicios internos. Finalmente, la capa de datos almacena entidades como usuarios, pacientes, médicos, historiales, tokens y registros de auditoría mediante un modelo relacional (Fowler, 2002; Bass et al., 2021).",
        ),
        (
            "normal",
            "Esta separación por capas facilita el mantenimiento del prototipo porque cada parte del sistema cumple una responsabilidad definida. La interfaz no accede directamente a la base de datos, sino que se comunica con servicios del backend mediante endpoints HTTP. A su vez, el backend encapsula reglas de negocio en controladores y servicios, lo cual permite modificar la lógica del sistema sin alterar necesariamente la estructura visual. Esta organización también favorece futuras ampliaciones, como agregar nuevos roles, nuevos tipos de registros clínicos o mecanismos adicionales de seguridad.",
        ),
        ("Heading 3", "2.1.6 Usabilidad e interacción con el usuario"),
        (
            "normal",
            "La usabilidad es un aspecto crítico en sistemas de salud porque los usuarios pueden tener distintos niveles de experiencia tecnológica. Un sistema seguro pero difícil de utilizar puede provocar errores, abandono o uso incorrecto de la información. Por ello, las interfaces del prototipo se diseñan con formularios claros, mensajes de validación, paneles con indicadores, tablas de consulta y estados vacíos que orientan al usuario cuando todavía no existen registros. Estos elementos buscan reducir ambigüedad y mejorar la comprensión del flujo de interacción (Nielsen, 1994; Norman, 2013).",
        ),
        (
            "normal",
            "En el caso del paciente, la interfaz prioriza acciones como consultar su historial, generar tokens, revisar accesos y actualizar datos personales. En el caso del médico, el diseño debe facilitar la validación de un token, consulta del historial autorizado y registro de una consulta médica cuando el permiso lo permite. En el caso del administrador, la interfaz se orienta a la supervisión de usuarios y validación de médicos. Esta diferenciación por rol ayuda a que cada usuario visualice únicamente las funciones necesarias para su operación dentro del sistema.",
        ),
        ("Heading 3", "2.2 Trabajos relacionados"),
    ]
    # Insert before first related heading while preserving the list order.
    ref = first_related
    for style, text in theoretical_items:
        p = OxmlElement("w:p")
        ref._element.addprevious(p)
        new_para = Paragraph(p, ref._parent)
        new_para.style = style
        run = new_para.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(12)
        if style.startswith("Heading"):
            run.bold = True
        new_para.paragraph_format.space_after = Pt(6)

# Fix related-work labels.
for p in doc.paragraphs:
    if p.text.strip().startswith("2.2.5. Lugo"):
        set_text(p, "2.2.6 Lugo Franco y Márquez Verdugo (2026)", "Heading 3")

for p in doc.paragraphs:
    if p.text.strip().startswith("2.3. Tabla comparativa"):
        set_text(p, "2.3 Tabla comparativa", "Heading 3")
        paragraph_after(
            p,
            "Al hacer la comparación se observa que las investigaciones revisadas abordan problemáticas similares desde enfoques parciales, concentrándose principalmente en aspectos como el acceso del paciente a su información, la interoperabilidad entre sistemas o la seguridad mediante tecnologías avanzadas. Algunas propuestas destacan la importancia del control del paciente sobre sus datos clínicos, mientras que otras priorizan mecanismos robustos de protección mediante criptografía o blockchain. Sin embargo, estas soluciones tienden a enfatizar un componente específico y no siempre integran, en un mismo prototipo web, historial médico, autorización temporal, niveles de acceso, roles diferenciados y auditoría operativa.",
        )
        break

# Cases of use section: correct the narrative so it matches the implemented system.
for p in doc.paragraphs:
    t = p.text.strip()
    replacements = {
        "En la figura 2 se presenta el diagrama general de casos de uso del Sistema de Gestión de Citas Médicas, el cual muestra la interacción entre los diferentes tipos de usuarios del sistema y las funcionalidades principales disponibles dentro de la plataforma.": "En la figura 2 se presenta el diagrama general de casos de uso del sistema web para la gestión segura del historial médico electrónico. El diagrama muestra la interacción entre los actores principales y las funcionalidades que permiten registrar usuarios, autenticar sesiones, consultar información clínica, generar autorizaciones temporales y mantener trazabilidad de los accesos.",
        "El sistema contempla tres actores principales: Paciente, Médico y Administrador, cada uno con diferentes permisos y responsabilidades dentro de la aplicación. Estos actores interactúan con el sistema para realizar diversas operaciones relacionadas con la gestión de citas médicas y la administración de usuarios.": "El sistema contempla tres actores principales: Paciente, Médico y Administrador. El paciente administra su perfil, consulta su historial, genera tokens y revisa auditoría; el médico consulta historiales únicamente cuando cuenta con autorización válida y registra nuevas consultas cuando el permiso lo permite; el administrador supervisa usuarios, valida médicos y mantiene el control general de las cuentas.",
        "El paciente es el usuario encargado de solicitar y gestionar sus citas médicas. Entre sus principales casos de uso se encuentran registrarse en el sistema, iniciar sesión, consultar la disponibilidad de médicos, agendar una cita médica, visualizar su historial de citas, cancelar citas previamente programadas y actualizar su información personal. Estas funcionalidades permiten que el paciente tenga control sobre la programación de sus consultas de manera sencilla y organizada.": "El paciente es el actor central de la propuesta, ya que conserva control sobre la información médica asociada a su cuenta. Entre sus principales casos de uso se encuentran registrarse, iniciar sesión, actualizar su perfil médico base, consultar su historial clínico, generar tokens de acceso, revisar sus tokens activos o expirados, revocar autorizaciones y consultar los eventos de auditoría vinculados con su expediente.",
        "Por otra parte, el médico interactúa con el sistema principalmente para gestionar su agenda y las citas asignadas. Entre sus funciones se encuentran iniciar sesión, visualizar su calendario de citas, consultar la información de los pacientes que atenderá, confirmar citas programadas, reprogramarlas en caso necesario y marcar las consultas como completadas una vez finalizada la atención médica. Además, el médico puede gestionar su disponibilidad de horarios dentro del sistema.": "El médico interactúa con el sistema como usuario clínico autorizado. Sus funciones se concentran en iniciar sesión, validar un token proporcionado por el paciente, consultar el historial médico permitido y registrar una consulta médica cuando el nivel de acceso otorgado lo autoriza. Además, para proteger la información sensible, el médico debe contar con una cuenta activa y validada antes de ejecutar operaciones clínicas dentro del sistema.",
        "Finalmente, el administrador tiene un rol de supervisión y gestión dentro de la plataforma. Este actor es responsable de administrar los usuarios del sistema, validar y aprobar el registro de nuevos médicos antes de que puedan utilizar completamente la plataforma, gestionar especialidades médicas, supervisar las citas registradas y mantener el control general del funcionamiento del sistema.": "Finalmente, el administrador tiene un rol de supervisión dentro de la plataforma. Este actor puede consultar usuarios registrados, cambiar roles, activar o desactivar cuentas y aprobar o rechazar médicos pendientes. Estas funciones permiten mantener un control mínimo sobre la identidad de los profesionales de salud y reducen el riesgo de que usuarios no validados accedan a funciones clínicas.",
        "En conjunto, el diagrama permite visualizar de forma clara cómo interactúan los diferentes actores con el sistema, así como las principales funcionalidades disponibles para cada tipo de usuario, proporcionando una visión general del comportamiento y alcance del sistema propuesto.": "En conjunto, los casos de uso permiten describir el comportamiento funcional esperado del sistema y sirven como puente entre los requerimientos, el diseño de la base de datos y la implementación de los módulos del prototipo. Las tablas siguientes presentan los casos de uso que se consideran centrales para la versión implementada y para las funciones que se proyectan completar dentro del mismo alcance.",
    }
    if t in replacements:
        set_text(p, replacements[t])

# Update duplicate CU-05 table to CU-06 and align it with the implemented audit module.
if len(doc.tables) > 9:
    t = doc.tables[9]
    set_table_field(t, "No. y nombre de CU:", "CU-06 – Consultar auditoría de accesos")
    set_table_field(t, "Fuentes:", "Requerimientos de trazabilidad, seguridad de la información y transparencia para el paciente.")
    set_table_field(t, "Actor(es):", "Paciente / Administrador")
    set_table_field(t, "Descripción:", "Permite consultar los registros de auditoría generados por el sistema, incluyendo acciones relacionadas con visualización de perfil, actualización de datos, generación de tokens, consulta de historial y registro de consultas médicas.")
    set_table_field(t, "Pre-Condiciones:", "El actor debe estar autenticado. / Deben existir eventos registrados o el sistema debe mostrar un estado informativo vacío.")
    set_table_field(t, "Flujo básico (escenario de éxito):", "El actor accede al módulo de auditoría. / El sistema consulta los registros asociados al usuario o al paciente. / El sistema muestra fecha, acción realizada, descripción e información técnica disponible. / El actor revisa los eventos para identificar accesos o modificaciones relevantes.")
    set_table_field(t, "Flujo alternativo (escenarios de fracaso):", "A1. No existen registros disponibles / El sistema muestra un mensaje informativo. / A2. Usuario sin permisos / El sistema bloquea la consulta y deniega el acceso.")
    set_table_field(t, "Post-Condiciones:", "El actor visualiza los eventos registrados sin modificar la información clínica.")
    set_table_field(t, "Notas:", "La auditoría no sustituye controles preventivos, pero fortalece la trazabilidad y permite detectar accesos inusuales.")

if len(doc.tables) > 10:
    t = doc.tables[10]
    set_table_field(t, "No. y nombre de CU:", "CU-07 – Registrar consulta médica")
    set_table_field(t, "Actor(es):", "Médico validado")
    set_table_field(t, "Descripción:", "Permite al médico validado registrar una nueva consulta médica en el historial del paciente, incluyendo motivo de consulta, diagnóstico, tratamiento y observaciones.")
    set_table_field(t, "Notas:", "En la implementación actual el backend valida que el médico exista y esté aprobado por el administrador. La restricción por nivel de token debe mantenerse como regla funcional del flujo de autorización.")

# Add two use cases as normal text tables after the last current table.
last_table = doc.tables[-1]
after = Paragraph(OxmlElement("w:p"), last_table._parent)
last_table._element.addnext(after._element)
after.add_run("Casos de uso complementarios propuestos para completar el alcance funcional:").bold = True

cu08 = doc.add_table(rows=0, cols=2)
for label, value in [
    ("No. y nombre de CU:", "CU-08 – Actualizar perfil médico base del paciente"),
    ("Fuentes:", "Requerimientos de gestión de perfil, disponibilidad de información básica y control del paciente sobre sus datos."),
    ("Actor(es):", "Paciente"),
    ("Descripción:", "Permite al paciente consultar y actualizar datos personales y clínicos básicos, como CURP, fecha de nacimiento, sexo, teléfono, dirección, grupo sanguíneo y alergias, según los campos definidos para el perfil."),
    ("Pre-Condiciones:", "Paciente autenticado. / El perfil de paciente debe existir en el sistema."),
    ("Flujo básico (escenario de éxito):", "El paciente accede a Mi Perfil Médico Base. / El sistema carga los datos existentes. / El paciente modifica los campos permitidos. / El sistema valida el formato de los datos. / El sistema guarda los cambios. / El sistema registra la acción en auditoría."),
    ("Flujo alternativo (escenarios de fracaso):", "A1. Datos inválidos / El sistema solicita corrección. / A2. Perfil no encontrado / El sistema informa que no existe perfil asociado. / A3. Error de almacenamiento / El sistema notifica la falla."),
    ("Post-Condiciones:", "El perfil del paciente queda actualizado y se registra el evento correspondiente."),
    ("Notas:", "El nombre de usuario puede mantenerse restringido para evitar inconsistencias de identidad; los datos clínicos básicos pueden ampliarse conforme avance el prototipo."),
]:
    row = cu08.add_row()
    row.cells[0].text = label
    row.cells[1].text = value
format_table(cu08)
after._element.addnext(cu08._element)

spacer = Paragraph(OxmlElement("w:p"), last_table._parent)
cu08._element.addnext(spacer._element)

cu09 = doc.add_table(rows=0, cols=2)
for label, value in [
    ("No. y nombre de CU:", "CU-09 – Administrar usuarios y validar médicos"),
    ("Fuentes:", "Requerimientos de control administrativo, gestión de identidad y seguridad de acceso a funciones clínicas."),
    ("Actor(es):", "Administrador"),
    ("Descripción:", "Permite al administrador consultar usuarios registrados, filtrar por rol, cambiar roles, activar o desactivar cuentas y aprobar o rechazar solicitudes de médicos antes de permitirles ejecutar operaciones clínicas."),
    ("Pre-Condiciones:", "Administrador autenticado. / El usuario objetivo debe existir en el sistema."),
    ("Flujo básico (escenario de éxito):", "El administrador accede al módulo de usuarios. / El sistema muestra la lista de cuentas registradas. / El administrador selecciona un usuario. / El administrador ejecuta una acción de aprobación, rechazo, cambio de rol o cambio de estado. / El sistema actualiza la cuenta y confirma la operación."),
    ("Flujo alternativo (escenarios de fracaso):", "A1. Usuario inexistente / El sistema informa que no se encontró el registro. / A2. Acción no permitida / El sistema bloquea la operación. / A3. Error de actualización / El sistema notifica la falla."),
    ("Post-Condiciones:", "El estado o rol del usuario queda actualizado conforme a la acción administrativa realizada."),
    ("Notas:", "La validación del médico es importante porque evita que una cuenta recién registrada acceda inmediatamente a operaciones de consulta o registro de información clínica."),
]:
    row = cu09.add_row()
    row.cells[0].text = label
    row.cells[1].text = value
format_table(cu09)
spacer._element.addnext(cu09._element)

# Make ER paragraphs normal rather than headings and correct figure references/entities.
for p in doc.paragraphs:
    if p.text.startswith("En la figura 2 se presenta el diagrama que corresponde"):
        set_text(p, "En la figura 3 se presenta el diagrama que corresponde al modelo entidad-relación lógico del sistema web propuesto para la gestión segura del historial médico electrónico centrado en el paciente. En él se representan las principales entidades del sistema, como Usuario, Paciente, Médico, Historial Médico, Token y Registro de Auditoría, así como las relaciones necesarias para asociar cada registro clínico con el paciente correspondiente y, cuando aplique, con el médico que lo generó.", "normal")
    elif p.text.startswith("El modelo separa la autenticación"):
        set_text(p, "El modelo separa la autenticación general del usuario de los perfiles clínicos específicos. La entidad Usuario concentra credenciales, rol y estado de la cuenta; Paciente y Médico almacenan información particular de cada tipo de usuario; Historial Médico representa las consultas registradas; Token modela las autorizaciones temporales emitidas por el paciente; y Registro de Auditoría conserva evidencia de acciones relevantes realizadas dentro de la plataforma.", "normal")
    elif p.text.startswith("Este modelo lógico sienta"):
        set_text(p, "Este modelo lógico sienta las bases para la implementación técnica del sistema en una base de datos relacional, favoreciendo integridad referencial, separación de responsabilidades y control sobre las operaciones que afectan información médica sensible.", "normal")

# Add implementation section content after 3.4 placeholder.
impl = find_para(doc, lambda p: p.text.strip() == "3.4 Descripción del sistema")
if impl:
    set_text(impl, "3.4 Descripción del sistema", "Heading 3")
    insert_block_after(
        impl,
        [
            ("normal", "El prototipo implementado corresponde a una aplicación web compuesta por un frontend desarrollado en Angular y un backend desarrollado con NestJS. La comunicación entre ambas capas se realiza mediante servicios HTTP que exponen operaciones para autenticación, pacientes, médicos, historiales médicos, tokens y auditoría. Esta estructura permite que la interfaz de usuario se mantenga separada de la lógica de negocio y que las reglas de seguridad se concentren principalmente en el servidor."),
            ("normal", "En el módulo de autenticación, el sistema permite registrar usuarios e iniciar sesión mediante correo electrónico y contraseña. Las contraseñas se almacenan cifradas y, al autenticarse correctamente, el usuario recibe un token JWT que se utiliza para identificar la sesión en solicitudes posteriores. Además, el sistema maneja roles de usuario, principalmente paciente, médico y administrador, lo cual permite restringir funcionalidades según el perfil correspondiente."),
            ("normal", "El módulo de pacientes permite consultar y actualizar el perfil asociado al usuario autenticado. Este perfil concentra datos personales y clínicos básicos que pueden apoyar la identificación del paciente dentro del sistema. El módulo de historial médico permite que el paciente consulte sus registros clínicos y que el médico validado pueda registrar nuevas consultas asociadas a un paciente. Cada registro médico almacena datos como fecha de consulta, motivo, diagnóstico, tratamiento y observaciones."),
            ("normal", "El módulo de tokens permite que el paciente genere códigos alfanuméricos únicos con vigencia temporal, nivel de acceso y descripción opcional. Estos tokens funcionan como mecanismo de autorización explícita para que un profesional de la salud pueda acceder al historial del paciente. En la implementación actual, el token se valida en el backend revisando existencia y fecha de expiración; como mejora funcional, debe completarse la revocación desde el servidor y reforzarse la validación del nivel de acceso antes de permitir operaciones de edición."),
            ("normal", "El módulo de auditoría registra eventos relevantes del sistema, como consulta de perfil, actualización de datos y creación de registros médicos. Esta bitácora permite al paciente revisar actividad relacionada con su expediente y proporciona una base para supervisar accesos. Aunque el prototipo se mantiene en un alcance académico, la inclusión de auditoría resulta fundamental porque conecta los requerimientos de seguridad con evidencia verificable dentro del sistema."),
        ],
    )

# Add relevant references missing from related-work table.
new_refs = [
    "Cichosz, S. L., Stausholm, M. N., Kronborg, T., Vestergaard, P., & Hejlesen, O. (2019). How to use blockchain for diabetes health care data and access management: An operational concept. Journal of Diabetes Science and Technology, 13(2), 248-253. https://doi.org/10.1177/1932296818790281",
    "Paydar, S., Emami, H., Asadi, F., Moghaddasi, H., & Hosseini, A. (2021). Functions and outcomes of personal health records for patients with chronic diseases: A systematic review. Perspectives in Health Information Management, 18(Spring), 1l.",
    "Zhan, W., Chen, C. L., Weng, W., Tsaur, W. J., Lim, Z. Y., & Deng, Y. Y. (2022). Incentive EMR sharing system based on consortium blockchain and IPFS. Healthcare, 10(10), 1840. https://doi.org/10.3390/healthcare10101840",
]
bib = find_para(doc, lambda p: p.text.strip().startswith("Bibliograf"))
anexos = find_para(doc, lambda p: p.text.strip() == "Anexos")
if bib and anexos:
    existing = "\n".join(p.text for p in doc.paragraphs)
    for ref in reversed(new_refs):
        if ref not in existing:
            p = OxmlElement("w:p")
            anexos._element.addprevious(p)
            para = Paragraph(p, anexos._parent)
            para.style = "normal"
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.first_line_indent = Inches(-0.5)
            para.paragraph_format.space_after = Pt(6)
            r = para.add_run(ref)
            r.font.name = "Arial"
            r.font.size = Pt(10)

doc.save(DOCX)
print(f"Expanded {DOCX}")
